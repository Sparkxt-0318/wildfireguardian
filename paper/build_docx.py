#!/usr/bin/env python
"""Build paper/WildfireGuardian_Park_2026.docx from paper/manuscript.md.

Markdown subset: see paper/README.md. No pandoc; python-docx only, so the build
runs anywhere the venv runs. Title page, numbered figures and tables, citations
numbered by first appearance, References generated from references.bib.

    python paper/build_docx.py [--md paper/manuscript.md] [--out paper/X.docx]
"""
from __future__ import annotations

import argparse
import datetime as dt
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
PAPER = REPO / "paper"
AUTHOR = "Siyeong Park (박시영)"
AFFIL = "Shanghai American School Puxi"


def parse_bib(path: Path) -> dict:
    """Minimal BibTeX reader: key -> dict(fields)."""
    txt = path.read_text(encoding="utf-8")
    out = {}
    for m in re.finditer(r"@(\w+)\s*\{\s*([^,\s]+)\s*,(.*?)\n\}", txt, flags=re.S):
        kind, key, body = m.groups()
        fields = {}
        for fm in re.finditer(r"(\w+)\s*=\s*\{((?:[^{}]|\{[^{}]*\})*)\}", body, flags=re.S):
            fields[fm.group(1).lower()] = re.sub(r"\s+", " ", fm.group(2)).strip()
        fields["_type"] = kind
        out[key] = fields
    return out


def fmt_ref(n: int, f: dict) -> str:
    authors = f.get("author", "").replace("{", "").replace("}", "")
    parts = [f"[{n}] {authors}" if authors else f"[{n}]", f'"{f.get("title", "")}."']
    for k in ("journal", "booktitle", "howpublished"):
        if f.get(k):
            parts.append(f.get(k) + ",")
    if f.get("year"):
        parts.append(f.get("year") + ".")
    if f.get("url"):
        parts.append(f.get("url"))
    return " ".join(parts)


def add_runs(par, text: str, cite_map: dict, gaps: list):
    """Inline markup: **bold**, *italic*, `code`, [@key], [GAP: ...]."""
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[@[^\]]+\]|\[GAP:[^\]]+\])")
    for m in pattern.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif tok.startswith("*"):
            par.add_run(tok[1:-1]).italic = True
        elif tok.startswith("[@"):
            keys = [k.strip().lstrip("@") for k in tok[2:-1].split(";")]
            nums = []
            for k in keys:
                if k not in cite_map:
                    cite_map[k] = len(cite_map) + 1
                nums.append(str(cite_map[k]))
            par.add_run("[" + ", ".join(nums) + "]")
        else:
            gaps.append(tok[5:-1].strip())
            r = par.add_run("[GAP: " + tok[5:-1].strip() + "]"); r.font.color.rgb = RGBColor(0xB4, 0x23, 0x18); r.italic = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def build(md_path: Path, out: Path) -> dict:
    bib = parse_bib(PAPER / "references.bib")
    lines = md_path.read_text(encoding="utf-8").split("\n")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, side, Cm(2.3))
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.15
    for name, size in (("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11.5)):
        h = doc.styles[name]; h.font.name = "Calibri"; h.font.size = Pt(size); h.font.color.rgb = RGBColor(0x1B, 0x1F, 0x23); h.font.bold = True

    cite_map: dict = {}
    gaps: list = []
    fig_n = tab_n = 0
    body_words = 0
    title = None
    i = 0
    in_list = None
    table_caption = None
    title_done = False

    def flush_table(rows):
        nonlocal tab_n
        tab_n += 1
        cap = doc.add_paragraph()
        cap.add_run(f"Table {tab_n}. ").bold = True
        add_runs(cap, table_caption or "", cite_map, gaps)
        t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Light Grid Accent 1"
        for r, row in enumerate(rows):
            for c, cell in enumerate(row):
                t.cell(r, c).text = ""
                add_runs(t.cell(r, c).paragraphs[0], cell, cite_map, gaps)
                for p in t.cell(r, c).paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9.5)
                        if r == 0:
                            run.bold = True
        doc.add_paragraph()

    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("# ") and not title_done:
            title = line[2:].strip(); title_done = True
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(120)
            r = p.add_run(title); r.bold = True; r.font.size = Pt(20)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(AUTHOR).font.size = Pt(13)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(AFFIL).font.size = Pt(11)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"Manuscript draft, built {dt.datetime.now(dt.timezone.utc).strftime('%Y-%m-%d')} from the repository's committed artifacts. Not for submission before the Korea Code Fair awards (December 2026)."); r.italic = True; r.font.size = Pt(9.5)
            doc.add_page_break()
            i += 1; continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1); i += 1; continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2); i += 1; continue
        m = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", line)
        if m:
            cap, rel = m.groups()
            img = (PAPER / rel) if not rel.startswith("/") else Path(rel)
            fig_n += 1
            if img.exists():
                doc.add_picture(str(img), width=Cm(16.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                gaps.append(f"figure file missing: {rel}")
                p = doc.add_paragraph(); r = p.add_run(f"[GAP: figure file missing: {rel}]"); r.font.color.rgb = RGBColor(0xB4, 0x23, 0x18)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Figure {fig_n}. ").bold = True
            add_runs(p, cap, cite_map, gaps)
            for run in p.runs:
                run.font.size = Pt(9.5)
            i += 1; continue
        m = re.match(r"^Table\s+\d*\.?\s*(.*)$", line)
        if m and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_caption = m.group(1).strip(); i += 1; continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.match(r"^:?-+:?$", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                flush_table(rows)
            table_caption = None
            continue
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            style_name = "List Number" if m.group(2)[0].isdigit() else "List Bullet"
            p = doc.add_paragraph(style=style_name); add_runs(p, m.group(3), cite_map, gaps)
            body_words += len(m.group(3).split()); i += 1; continue
        if line.strip() == "":
            i += 1; continue
        # paragraph: join following non-empty, non-special lines
        buf = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not re.match(r"^(#|!\[|\||Table\s|\s*[-*]\s|\s*\d+\.\s)", lines[j]):
            buf.append(lines[j].rstrip()); j += 1
        text = " ".join(buf)
        p = doc.add_paragraph(); add_runs(p, text, cite_map, gaps)
        body_words += len(re.sub(r"\[@[^\]]+\]", "", text).split())
        i = j

    # references
    if cite_map:
        doc.add_heading("References", level=1)
        missing = [k for k in cite_map if k not in bib]
        for k, n in sorted(cite_map.items(), key=lambda kv: kv[1]):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.first_line_indent = Cm(-0.8)
            p.add_run(fmt_ref(n, bib[k]) if k in bib else f"[{n}] MISSING BIB ENTRY: {k}").font.size = Pt(9.5)
    else:
        missing = []
    doc.save(out)
    return {"title": title, "body_words": body_words, "figures": fig_n, "tables": tab_n,
            "references": len(cite_map), "missing_refs": missing, "gaps": gaps, "out": str(out)}


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--md", default=str(PAPER / "manuscript.md"))
    ap.add_argument("--out", default=str(PAPER / "WildfireGuardian_Park_2026.docx"))
    a = ap.parse_args()
    info = build(Path(a.md), Path(a.out))
    print("[paper] " + json.dumps({k: v for k, v in info.items() if k != "gaps"}) + f" gaps={len(info['gaps'])}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
